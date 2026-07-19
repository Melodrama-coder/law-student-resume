#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_resume.py - 法学生求职简历生成器

输入: 一份描述简历内容的 JSON 文件
输出: 一页 A4 的 Word(.docx) 与 PDF(.pdf)

特性:
  - 多模板: classic / minimal / centered / modern (版式风格各异)
  - 字体可选: 仿宋 / 微软雅黑 / 宋体 / 楷体 (中文); 英文统一 Times New Roman
  - 自动适配: 若内容超出一页, 自动缩小字号直至正好一页
  - 风格: 简洁、无花哨、清晰大方
  - 数据真实性: 脚本只负责排版, 内容完全来自用户提供的 JSON (绝不编造)

用法:
  python generate_resume.py --data resume.json --out ./output [--font yahei] [--template classic]
"""
import argparse
import json
import os
import sys
import datetime
import logging
from exceptions import ColorParseError, DataValidationError, TemplateNotFoundError
from __version__ import __version__

# ---------------------------------------------------------------------------
# 日志: 同时输出到控制台与 resume.log (便于排错, 见验证命令 cat resume.log)
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("resume.log", encoding="utf-8"),
        logging.StreamHandler(),
    ],
)
logger = logging.getLogger(__name__)


def _friendly_import_error(missing):
    """依赖缺失时给出友好安装提示，而非抛出堆栈。"""
    print("=" * 60)
    print("  缺少必要的 Python 依赖：%s" % missing)
    print("  请先安装依赖后重试：")
    print("    pip install python-docx pypdf pillow")
    print("  （生成 PDF 还需安装 Microsoft Word 或 LibreOffice；")
    print("   岗位匹配雷达图另需 pip install matplotlib）")
    print("=" * 60)
    sys.exit(1)


try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as _e:
    _friendly_import_error("python-docx (%s)" % _e)

# ---------------------------------------------------------------------------
# 字体映射
# ---------------------------------------------------------------------------
CN_FONT_MAP = {
    "fang": "仿宋",        # 仿宋 (FangSong)
    "yahei": "微软雅黑",   # Microsoft YaHei
    "song": "宋体",        # SimSun
    "kaiti": "楷体",       # KaiTi
}
EN_FONT = "Times New Roman"
DARK = RGBColor(0x22, 0x22, 0x22)
MIDGRAY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ---------------------------------------------------------------------------
# 配色方案 (强调色可命名选取, 也可传入 #RRGGBB 自定义)
# 主推: 黑 / 白 / 灰 + 天蓝色; 其余为“想要其他颜色”时的备选。
# ---------------------------------------------------------------------------
COLOR_PRESETS = {
    # 预设名与色值严格对应 SKILL.md §13.5
    "skyblue":  (135, 206, 235),   # 天蓝 #87CEEB
    "black":    (0, 0, 0),         # 纯黑 #000000
    "gray":     (128, 128, 128),   # 灰   #808080
    "darkgray": (64, 64, 64),      # 深灰 #404040
    "navy":     (0, 0, 128),       # 藏青 #000080
    "teal":     (0, 128, 128),     # 墨绿 #008080
    "burgundy": (128, 0, 32),      # 酒红 #800020
    "green":    (0, 128, 0),       # 森林绿 #008000
    "purple":   (128, 0, 128),     # 紫色 #800080
}
DEFAULT_ACCENT = "skyblue"


def parse_color(color_input):
    """解析颜色输入，支持预设名或十六进制，返回 RGB 元组 (r, g, b)。

    预设名: skyblue / black / gray / darkgray / navy / teal / burgundy / green / purple
    十六进制: #RRGGBB 或 #RGB
    见 SKILL.md §13.4。无效输入抛出 ColorParseError。
    """
    PRESET = COLOR_PRESETS
    if not color_input or not isinstance(color_input, str):
        raise ColorParseError(
            "颜色值不能为空，请使用预设名 %s 或 #RRGGBB 格式" % list(PRESET.keys()))
    cleaned = color_input.strip().lower()
    if cleaned in PRESET:
        return PRESET[cleaned]
    if cleaned.startswith("#"):
        hex_str = cleaned.lstrip("#")
        if len(hex_str) == 3:
            hex_str = "".join(c * 2 for c in hex_str)
        if len(hex_str) == 6:
            try:
                return tuple(int(hex_str[i:i + 2], 16) for i in (0, 2, 4))
            except ValueError:
                raise ColorParseError(
                    '无效的十六进制颜色值: "%s"，请使用 #RRGGBB 格式' % color_input)
    raise ColorParseError(
        '不支持的颜色值: "%s"。请使用预设名 %s 或 #RRGGBB 格式。'
        % (color_input, list(PRESET.keys())))


def resolve_color(spec, default=DEFAULT_ACCENT):
    """向后兼容封装: 空值回落到 default，其余委托 parse_color（无效则抛 ColorParseError）。"""
    if not spec:
        spec = default
    return parse_color(spec)


def lighten(rgb, factor=0.55):
    """把颜色向白色混合, 用于生成浅色分隔线。factor 越大越浅。"""
    return tuple(int(c + (255 - c) * factor) for c in rgb)


def color_brightness(rgb):
    """计算颜色感知亮度 (0-255)，sRGB 加权。见 §13.2。"""
    if isinstance(rgb, RGBColor):
        rgb = (rgb[0], rgb[1], rgb[2])
    return 0.299 * rgb[0] + 0.587 * rgb[1] + 0.114 * rgb[2]


def get_contrast_text_color(rgb):
    """根据背景颜色自动选择文字色: 亮度 > 160 → 黑色，否则白色。返回 (r, g, b) 元组。见 §13.2。"""
    if isinstance(rgb, RGBColor):
        rgb = (rgb[0], rgb[1], rgb[2])
    return (0, 0, 0) if color_brightness(rgb) > 160 else (255, 255, 255)


def contrast_text(accent_rgb, force_color=None):
    """根据强调色亮度自动返回可读文字颜色 (RGBColor)，用于填充标题栏/色块背景。

    浅色背景 → 黑字；深色背景 → 白字（亮度阈值见 §13.2）。
    force_color 非空时直接返回该颜色（如 --force-color 覆盖，不推荐）。
    """
    if force_color is not None:
        return force_color
    return RGBColor(*get_contrast_text_color(accent_rgb))



# ---------------------------------------------------------------------------
# 模板预设: 控制强调色、页眉风格、章节标题风格、是否显示英文副标题
#   accent : 默认强调色 (预设名), 可被 meta.accent / --accent 覆盖
#   rule   : 分隔线颜色模式 -> "accent"(同强调色) / "light"(浅化强调色) / (r,g,b)固定
# ---------------------------------------------------------------------------
TEMPLATES = {
    "classic": {
        "accent": "skyblue", "rule": "accent",
        "en_sub": True, "title": "underline", "header": "left",
        "desc": "经典天蓝 · 标题带英文副标题+分隔线 · 适合红圈所/综合投递",
    },
    "minimal": {
        "accent": "black", "rule": (0xBB, 0xBB, 0xBB),
        "en_sub": False, "title": "rule", "header": "left",
        "desc": "极简黑白 · 无彩色、细灰线 · 适合央企法务/公检法等稳重场合",
    },
    "centered": {
        "accent": "darkgray", "rule": "accent",
        "en_sub": False, "title": "center", "header": "center",
        "desc": "居中庄重 · 黑灰稳重、姓名/标题居中 · 配仿宋更显体制内正式风格",
    },
    "modern": {
        "accent": "skyblue", "rule": "light",
        "en_sub": True, "title": "bar", "header": "bar",
        "desc": "现代天蓝 · 姓名色块+方块标题 · 适合新锐律所/互联网法务",
    },
    "compact": {
        "accent": "black", "rule": "black",
        "en_sub": False, "title": "fill", "header": "compact",
        "desc": "紧凑填满 · 黑色填充标题栏+右上证件照+编号bullet · 一页饱满",
    },
}


def resolve_palette(T, accent_override=None):
    """根据模板 + 可选覆盖色, 计算 (accent, rule) 两个 RGBColor。"""
    accent_rgb = resolve_color(accent_override if accent_override else T.get("accent"))
    rmode = T.get("rule", "accent")
    if isinstance(rmode, (tuple, list)):
        rule_rgb = tuple(rmode)
    elif rmode == "light":
        rule_rgb = lighten(accent_rgb, 0.55)
    else:  # "accent"
        rule_rgb = accent_rgb
    return RGBColor(*accent_rgb), RGBColor(*rule_rgb)

# 章节顺序 (核心在前, 可选项按需出现)
SECTION_ORDER = [
    ("education",    "教育经历",    "Education"),
    ("work",         "工作经历",    "Work Experience"),
    ("internship",   "实习经历",    "Internship"),
    ("research",     "科研经历",    "Research"),
    ("competition",  "竞赛经历",    "Competitions"),
    ("project",      "项目经历",    "Projects"),
    ("skills",       "专业技能",    "Skills"),
    ("summary",      "个人总结",    "Summary"),
]
# 章节 key -> (中文标题, 英文副标题) 映射, 供自定义顺序/多版本复用
SECTION_TITLES = {k: (cn, en) for k, cn, en in SECTION_ORDER}

# ---------------------------------------------------------------------------
# 法律术语英译词典 (英文版简历用)。
#   - section : 章节标题的英文
#   - label   : 抬头/技能标签的英文
#   - degree  : 常见学位/资格的标准英译, 供 agent 组织 data['en'] 时参考
# 说明: 英文自由文本(经历 bullet)不做机器翻译, 由 agent 依据 references/legal_terms.md
#       翻译后写入 data['en'] 并行块; 脚本只负责把"固定骨架"翻成标准英文。
# ---------------------------------------------------------------------------
EN_SECTION = {
    "education":   "EDUCATION",
    "work":        "WORK EXPERIENCE",
    "internship":  "INTERNSHIP EXPERIENCE",
    "research":    "RESEARCH EXPERIENCE",
    "competition": "COMPETITIONS & MOOTS",
    "project":     "PROJECTS",
    "skills":      "SKILLS & QUALIFICATIONS",
    "summary":     "PROFILE",
}
EN_LABEL = {
    "phone":        "Tel",
    "email":        "Email",
    "qualification": "Qualification",
    "language":     "Languages",
    "software":     "IT & Legal Research",
}
# 学位 / 资格 标准英译 (供 references/legal_terms.md 与 agent 组织英文数据参考)
LEGAL_TERMS_HINT = {
    "法学硕士": "Master of Laws (LL.M.)",
    "法律硕士": "Juris Master (J.M.)",
    "法学学士": "Bachelor of Laws (LL.B.)",
    "法律职业资格": "PRC Legal Profession Qualification (Bar)",
    "证券从业资格": "Securities Qualification Certificate",
    "尽职调查": "due diligence",
    "法律意见书": "legal opinion",
    "尽调报告": "due diligence report",
    "合规审查": "compliance review",
    "商事仲裁": "commercial arbitration",
    "类案检索": "case-law research",
    "模拟法庭": "moot court",
}

# 多版本预设：一稿多投时按求职方向重排模块顺序。
#   order : 章节出现的顺序 (出现在前 = 更突出); 不存在的章节自动跳过。
#   label : 版本中文名, 用于在文件名与抬头标注。
VERSION_PRESETS = {
    "law":   {"label": "律所版",
              "order": ["internship", "work", "education", "skills",
                        "research", "competition", "project", "summary"]},
    "legal":  {"label": "法务版",
              "order": ["work", "internship", "education", "skills",
                        "project", "research", "competition", "summary"]},
    "civil":  {"label": "体制版",
              "order": ["education", "competition", "research",
                        "internship", "work", "skills", "summary", "project"]},
}


# ---------------------------------------------------------------------------
# 工具
# ---------------------------------------------------------------------------
def hexstr(color):
    try:
        return "%02X%02X%02X" % (color[0], color[1], color[2])
    except Exception:
        return "1F3A5F"


def set_run_font(run, cn_font, size, bold=False, color=DARK):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = EN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), EN_FONT)
    rfonts.set(qn("w:hAnsi"), EN_FONT)
    rfonts.set(qn("w:cs"), EN_FONT)
    rfonts.set(qn("w:eastAsia"), cn_font)


def set_para_margins(p, before=0, after=2, line=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if line is not None:
        p.paragraph_format.line_spacing = line


def add_bottom_border(p, color, size=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), hexstr(color))
    pbdr.append(bottom)
    pPr.append(pbdr)


def set_top_border(p, color, size=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size))
    top.set(qn("w:space"), "2")
    top.set(qn("w:color"), hexstr(color))
    pbdr.append(top)
    pPr.append(pbdr)


def set_para_shading(p, fill_hex):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def set_right_tab(p, position_cm):
    p.paragraph_format.tab_stops.add_tab_stop(Cm(position_cm), WD_TAB_ALIGNMENT.RIGHT)


def add_text(p, text, cn_font, size, bold=False, color=DARK, font_override=None):
    run = p.add_run(text)
    if font_override:
        run.font.name = font_override
    set_run_font(run, cn_font, size, bold, color)
    return run


def build_contact(pers, lang="zh"):
    contact = []
    if lang == "en":
        if pers.get("phone"):
            contact.append(f"{EN_LABEL['phone']}: {pers['phone']}")
        if pers.get("email"):
            contact.append(f"{EN_LABEL['email']}: {pers['email']}")
        # 英文版其余抬头字段直接使用 data['en'] 中已翻译好的自由文本
        for k in ("gender_age", "native_place", "political", "job_intention"):
            if pers.get(k):
                contact.append(pers[k])
        return contact
    if pers.get("phone"):
        contact.append(f"电话：{pers['phone']}")
    if pers.get("email"):
        contact.append(f"邮箱：{pers['email']}")
    if pers.get("gender_age"):
        contact.append(pers["gender_age"])
    if pers.get("native_place"):
        contact.append(pers["native_place"])
    if pers.get("political"):
        contact.append(pers["political"])
    if pers.get("job_intention"):
        contact.append(pers["job_intention"])
    return contact


# ---------------------------------------------------------------------------
# 页眉 (随模板变化)
# ---------------------------------------------------------------------------
def build_header(doc, data, cn_font, S, T, accent, rule, lang="zh", force_color=None):
    pers = data.get("personal", {})
    photo = pers.get("photo") or data.get("meta", {}).get("photo")
    name = pers.get("name", "姓名" if lang == "zh" else "Name")
    contact = build_contact(pers, lang)
    style = T["header"]

    # 带照片时统一用左右两栏表格 (照片置右)
    if photo and os.path.exists(photo):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        tbl = table._tbl
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "none")
            borders.append(e)
        tbl.tblPr.append(borders)
        left = table.cell(0, 0)
        right = table.cell(0, 1)
        left.width = Cm(14.5)
        right.width = Cm(3.2)
        p0 = left.paragraphs[0]
        set_para_margins(p0, before=0, after=1)
        add_text(p0, name, cn_font, S["name"], bold=True, color=accent if style != "left" else DARK)
        if contact:
            p1 = left.add_paragraph()
            set_para_margins(p1, before=0, after=1)
            add_text(p1, "  |  ".join(contact), cn_font, S["body"])
        try:
            right.paragraphs[0].add_run().add_picture(photo, width=Cm(3.0))
        except Exception:
            pass
        sep = doc.add_paragraph()
        set_para_margins(sep, before=1, after=1)
        add_bottom_border(sep, rule, 12)
        return

    # 无照片: 按模板风格
    if style == "bar":
        pn = doc.add_paragraph()
        set_para_shading(pn, hexstr(accent))
        set_para_margins(pn, before=0, after=1)
        pn.paragraph_format.left_indent = Cm(0.15)
        add_text(pn, "  " + name, cn_font, S["name"], bold=True,
                 color=contrast_text(accent, force_color))
        if contact:
            pc = doc.add_paragraph()
            set_para_margins(pc, before=1, after=1)
            add_text(pc, "  |  ".join(contact), cn_font, S["body"])
    elif style == "compact":
        # 紧凑布局: 姓名左上, 联系信息在下方一行; 照片在右侧 (单独表格)
        p0 = doc.add_paragraph()
        set_para_margins(p0, before=0, after=1)
        add_text(p0, name, cn_font, S["name"] + 1, bold=True)
        if contact:
            p1 = doc.add_paragraph()
            set_para_margins(p1, before=0, after=1)
            add_text(p1, "  |  ".join(contact), cn_font, S["body"])
    elif style == "center":
        pn = doc.add_paragraph()
        pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_margins(pn, before=0, after=1)
        add_text(pn, name, cn_font, S["name"], bold=True, color=accent)
        if contact:
            pc = doc.add_paragraph()
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_para_margins(pc, before=0, after=1)
            add_text(pc, "  |  ".join(contact), cn_font, S["body"])
        sep = doc.add_paragraph()
        set_para_margins(sep, before=1, after=1)
        add_bottom_border(sep, rule, 8)
    else:  # left (classic / minimal)
        p0 = doc.add_paragraph()
        set_para_margins(p0, before=0, after=1)
        add_text(p0, name, cn_font, S["name"], bold=True)
        if contact:
            p1 = doc.add_paragraph()
            set_para_margins(p1, before=0, after=1)
            add_text(p1, "  |  ".join(contact), cn_font, S["body"])
        sep = doc.add_paragraph()
        set_para_margins(sep, before=1, after=1)
        add_bottom_border(sep, rule, 12)


# ---------------------------------------------------------------------------
# 章节标题 (随模板变化)
# ---------------------------------------------------------------------------
def add_section_title(doc, title_cn, title_en, cn_font, S, T, accent, rule, force_color=None):
    style = T["title"]
    p = doc.add_paragraph()
    set_para_margins(p, before=2, after=1)

    if style == "bar":
        # 现代标题: 左=色块标记+中文, 右=英文副标题 (同基线, 右对齐 tab)
        set_right_tab(p, 17.0)
        add_text(p, "■", cn_font, S["head"] - 1, bold=True, color=accent)
        add_text(p, " " + title_cn, cn_font, S["head"], bold=True, color=accent)
        if T["en_sub"] and title_en:
            p.add_run("\t")
            run = p.add_run(title_en)
            set_run_font(run, cn_font, S["head"] - 1, False, MIDGRAY)
        add_bottom_border(p, rule, 4)
    elif style == "fill":
        set_para_shading(p, hexstr(accent))
        set_top_border(p, accent, 2)
        set_para_margins(p, before=3, after=2)
        add_text(p, "  " + title_cn, cn_font, S["head"], bold=True,
                 color=contrast_text(accent, force_color))
    elif style == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, title_cn, cn_font, S["head"], bold=True, color=accent)
        add_bottom_border(p, rule, 6)
    elif style == "rule":
        add_text(p, title_cn, cn_font, S["head"], bold=True, color=accent)
        add_bottom_border(p, rule, 6)
    else:  # underline (classic)
        add_text(p, title_cn, cn_font, S["head"], bold=True, color=accent)
        if T["en_sub"] and title_en:
            run = p.add_run(f"   {title_en}")
            set_run_font(run, cn_font, S["head"] - 2, False, MIDGRAY)
        add_bottom_border(p, rule, 6)
    return p


def add_two_line(doc, left_bold, right_text, sub_line, cn_font, S, center=False):
    p = doc.add_paragraph()
    set_para_margins(p, before=1.5, after=0)
    usable = 17.0
    set_right_tab(p, usable)
    add_text(p, left_bold, cn_font, S["entry"], bold=True)
    if right_text:
        p.add_run("\t").font.size = Pt(S["entry"])
        add_text(p, right_text, cn_font, S["entry"], color=MIDGRAY)
    if sub_line:
        p2 = doc.add_paragraph()
        set_para_margins(p2, before=0, after=0.5)
        add_text(p2, sub_line, cn_font, S["sub"], color=MIDGRAY)


def add_bullets(doc, bullets, cn_font, S, indent=0.6, numbered=False):
    for i, b in enumerate(bullets or []):
        if numbered:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(indent)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            set_para_margins(p, before=0, after=0.3, line=1.0)
            add_text(p, f"{i+1}. {b}", cn_font, S["body"])
        else:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(indent)
            p.paragraph_format.first_line_indent = Cm(-0.3)
            set_para_margins(p, before=0, after=0.3, line=1.0)
            add_text(p, b, cn_font, S["body"])


def add_wrapped_line(doc, text, cn_font, S, bold=False, color=DARK):
    p = doc.add_paragraph()
    set_para_margins(p, before=0, after=0.3, line=1.0)
    add_text(p, text, cn_font, S["body"], bold=bold, color=color)


def add_labeled_line(doc, text, cn_font, S, bold_color=DARK):
    """专业技能等模块：冒号前面的标签加粗，冒号及之后正文正常字重；无冒号则整行正常。
    同时兼容中文冒号『：』与英文冒号 ': '（英文版简历用）。"""
    p = doc.add_paragraph()
    set_para_margins(p, before=0, after=0.3, line=1.0)
    idx = -1
    clen = 1
    if "：" in text:
        idx = text.index("：")
        clen = 1
    elif ": " in text:
        idx = text.index(": ")
        clen = 2
    if idx >= 0:
        add_text(p, text[:idx + clen], cn_font, S["body"], bold=True, color=bold_color)
        rest = text[idx + clen:]
        if rest:
            add_text(p, rest, cn_font, S["body"], bold=False)
    else:
        add_text(p, text, cn_font, S["body"], bold=False)


# ---------------------------------------------------------------------------
# 主构建
# ---------------------------------------------------------------------------
def build_document(data, cn_font, base_size, T, accent_override=None,
                   section_order=None, lang="zh", force_color=None):
    accent, rule = resolve_palette(T, accent_override)

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)

    S = {
        "name": base_size + 5,
        "head": base_size + 1,
        "entry": base_size,
        "sub": base_size - 1,
        "body": base_size - 0.5,
    }
    for k in S:
        S[k] = max(S[k], 6.5)

    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = Pt(S["body"])
    style.paragraph_format.line_spacing = 1.0
    style.element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)

    build_header(doc, data, cn_font, S, T, accent, rule, lang, force_color)

    is_en = (lang == "en")

    def _title(key):
        """返回该章节的 (主标题, 副标题)。英文版主标题用英文, 不带副标题。"""
        cn, en = SECTION_TITLES[key]
        if is_en:
            return EN_SECTION.get(key, en or cn), ""
        return cn, en

    sections = section_order or [k for k, _, _ in SECTION_ORDER]
    for key in sections:
        if key not in SECTION_TITLES:
            continue
        title_cn, title_en = _title(key)
        if key == "skills":
            skills = data.get("skills")
            if not skills:
                continue
            add_section_title(doc, title_cn, title_en, cn_font, S, T, accent, rule, force_color)
            lines = []
            if is_en:
                if skills.get("qualification"):
                    lines.append(f"{EN_LABEL['qualification']}: {skills['qualification']}")
                if skills.get("language"):
                    lines.append(f"{EN_LABEL['language']}: {skills['language']}")
                if skills.get("software"):
                    lines.append(f"{EN_LABEL['software']}: {skills['software']}")
            else:
                if skills.get("qualification"):
                    lines.append(f"职业资格：{skills['qualification']}")
                if skills.get("language"):
                    lines.append(f"语言能力：{skills['language']}")
                if skills.get("software"):
                    lines.append(f"办公与检索：{skills['software']}")
            for extra in skills.get("extra", []) or []:
                lines.append(extra)
            for ln in lines:
                add_labeled_line(doc, ln, cn_font, S)
            continue

        if key == "summary":
            summary = data.get("summary")
            if not summary:
                continue
            add_section_title(doc, title_cn, title_en, cn_font, S, T, accent, rule, force_color)
            add_wrapped_line(doc, summary, cn_font, S)
            continue

        entries = data.get(key)
        if not entries:
            continue
        add_section_title(doc, title_cn, title_en, cn_font, S, T, accent, rule, force_color)

        for e in entries:
            if key in ("education",):
                left = e.get("school", "")
                right = e.get("dates", "")
                sub = "  ·  ".join(
                    x for x in [e.get("college"), e.get("major"), e.get("degree"), e.get("location")] if x
                )
                add_two_line(doc, left, right, sub, cn_font, S)
                detail = []
                if e.get("gpa"):
                    detail.append(e["gpa"])
                if e.get("rank"):
                    detail.append(e["rank"])
                if detail:
                    add_wrapped_line(doc, "  ".join(detail), cn_font, S, color=MIDGRAY)
                if e.get("courses"):
                    add_wrapped_line(doc, e["courses"], cn_font, S, color=MIDGRAY)
                if e.get("honors"):
                    add_wrapped_line(doc, e["honors"], cn_font, S, color=MIDGRAY)

            elif key in ("internship", "work", "research", "competition", "project"):
                left = e.get("org") or e.get("title") or e.get("name") or ""
                right = e.get("dates", "")
                role = e.get("role") or e.get("dept_role") or ""
                location = e.get("location", "")
                sub = "  ·  ".join(x for x in [role, location] if x)
                add_two_line(doc, left, right, sub, cn_font, S)
                desc = e.get("description")
                if desc:
                    if isinstance(desc, list):
                        add_bullets(doc, desc, cn_font, S, numbered=(T.get("title") == "fill"))
                    else:
                        add_wrapped_line(doc, desc, cn_font, S, color=MIDGRAY)
                if e.get("bullets"):
                    add_bullets(doc, e["bullets"], cn_font, S, numbered=(T.get("title") == "fill"))

    return doc


# ---------------------------------------------------------------------------
# PDF 转换 (Microsoft Word -> 失败回退 LibreOffice)
# ---------------------------------------------------------------------------
def convert_to_pdf(docx_path, pdf_path):
    if os.path.exists(pdf_path):
        try:
            os.remove(pdf_path)
        except Exception:
            pass
    try:
        import win32com.client as wc
        word = wc.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        doc.Close(False)
        word.Quit()
        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            return True, "Microsoft Word"
    except Exception as ex:
        print(f"  [warn] Word 转换失败: {ex}")
    try:
        import subprocess
        for so in ["soffice", "libreoffice"]:
            try:
                subprocess.run(
                    [so, "--headless", "--convert-to", "pdf", "--outdir",
                     os.path.dirname(pdf_path), os.path.abspath(docx_path)],
                    capture_output=True, timeout=120,
                )
            except Exception:
                continue
            if os.path.exists(pdf_path):
                return True, "LibreOffice"
    except Exception:
        pass
    return False, None


def count_pdf_pages(pdf_path):
    try:
        from pypdf import PdfReader
        return len(PdfReader(pdf_path).pages)
    except Exception:
        return 99


def stamp_pdf_metadata(pdf_path, title="Resume", author=""):
    """在 PDF 元数据写入生成时间戳与制作信息, 便于版本管理。"""
    try:
        from pypdf import PdfReader, PdfWriter
        now = datetime.datetime.now()
        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        d = f"D:{now.strftime('%Y%m%d%H%M%S')}"
        writer.add_metadata({
            "/Title": title,
            "/Author": author or title,
            "/Producer": "law-student-resume skill",
            "/Creator": "law-student-resume skill",
            "/CreationDate": d,
            "/ModDate": d,
            "/GeneratedAt": now.strftime("%Y-%m-%d %H:%M:%S"),
        })
        tmp = pdf_path + ".tmp"
        with open(tmp, "wb") as fh:
            writer.write(fh)
        os.replace(tmp, pdf_path)
        return now.strftime("%Y-%m-%d %H:%M:%S")
    except Exception as ex:
        print(f"  [warn] 写入 PDF 时间戳失败: {ex}")
        return None


# ---------------------------------------------------------------------------
# 语言解析: 英文版使用 data['en'] 并行块的内容 (meta 保留)
# ---------------------------------------------------------------------------
def resolve_lang_data(data, lang):
    """返回 (用于构建的 data, 英文内容是否齐备)。
    lang=='en' 时用 data['en'] 覆盖对应内容块; 缺失则回退中文并返回 False。"""
    if lang != "en":
        return data, True
    en = data.get("en")
    if not isinstance(en, dict) or not en:
        return data, False
    merged = dict(data)
    for k, v in en.items():
        merged[k] = v
    return merged, True


# ---------------------------------------------------------------------------
# LaTeX 导出 (隐藏功能 --format latex): 用模板字符串生成 .tex, 无需 pylatex
# ---------------------------------------------------------------------------
def _tex_escape(s):
    if s is None:
        return ""
    s = str(s)
    repl = {
        "\\": r"\textbackslash{}", "&": r"\&", "%": r"\%", "$": r"\$",
        "#": r"\#", "_": r"\_", "{": r"\{", "}": r"\}",
        "~": r"\textasciitilde{}", "^": r"\textasciicircum{}",
    }
    for a, b in repl.items():
        s = s.replace(a, b)
    return s


def export_latex(data, tex_path, lang="zh", section_order=None):
    """生成 LaTeX .tex 文件 (article 类, 依赖 ctex 处理中文)。"""
    d, _ = resolve_lang_data(data, lang)
    pers = d.get("personal", {})
    esc = _tex_escape
    is_en = (lang == "en")

    def title(key):
        cn, en = SECTION_TITLES.get(key, (key, key))
        return EN_SECTION.get(key, en) if is_en else cn

    lines = []
    lines.append(r"\documentclass[10.5pt,a4paper]{article}")
    if is_en:
        lines.append(r"\usepackage[T1]{fontenc}")
        lines.append(r"\usepackage{times}")
    else:
        lines.append(r"\usepackage{ctex}")
    lines.append(r"\usepackage[margin=1.4cm]{geometry}")
    lines.append(r"\usepackage{enumitem,titlesec,xcolor}")
    lines.append(r"\definecolor{accent}{HTML}{2E7CC2}")
    lines.append(r"\titleformat{\section}{\large\bfseries\color{accent}}"
                 r"{}{0em}{}[\titlerule]")
    lines.append(r"\titlespacing{\section}{0pt}{6pt}{3pt}")
    lines.append(r"\setlist[itemize]{leftmargin=1.4em,itemsep=1pt,topsep=1pt}")
    lines.append(r"\pagestyle{empty}")
    lines.append(r"\begin{document}")

    # 抬头
    lines.append(r"\begin{center}")
    lines.append(r"{\LARGE\bfseries " + esc(pers.get("name", "")) + r"}\\[2pt]")
    contact = build_contact(pers, lang)
    if contact:
        lines.append(esc("  |  ".join(contact)))
    lines.append(r"\end{center}")

    order = section_order or [k for k, _, _ in SECTION_ORDER]
    for key in order:
        if key == "skills":
            skills = d.get("skills")
            if not skills:
                continue
            lines.append(r"\section*{" + esc(title(key)) + "}")
            rows = []
            lbl = EN_LABEL if is_en else {
                "qualification": "职业资格", "language": "语言能力",
                "software": "办公与检索"}
            sep = ": " if is_en else "："
            for fk in ("qualification", "language", "software"):
                if skills.get(fk):
                    rows.append(r"\textbf{" + esc(lbl[fk]) + esc(sep) + "}" +
                                esc(skills[fk]))
            for extra in skills.get("extra", []) or []:
                rows.append(esc(extra))
            lines.append(r"\\".join(rows))
            continue
        if key == "summary":
            if not d.get("summary"):
                continue
            lines.append(r"\section*{" + esc(title(key)) + "}")
            lines.append(esc(d["summary"]))
            continue
        entries = d.get(key)
        if not entries:
            continue
        lines.append(r"\section*{" + esc(title(key)) + "}")
        for e in entries:
            if key == "education":
                head = esc(e.get("school", ""))
                right = esc(e.get("dates", ""))
                sub = "  ·  ".join(x for x in [e.get("college"), e.get("major"),
                                  e.get("degree"), e.get("location")] if x)
                lines.append(r"\textbf{" + head + r"} \hfill " + right + r"\\")
                if sub:
                    lines.append(esc(sub) + r"\\")
                for fk in ("gpa", "rank", "courses", "honors"):
                    if e.get(fk):
                        lines.append(esc(e[fk]) + r"\\")
            else:
                head = esc(e.get("org") or e.get("title") or e.get("name") or "")
                right = esc(e.get("dates", ""))
                role = e.get("role") or e.get("dept_role") or ""
                loc = e.get("location", "")
                sub = "  ·  ".join(x for x in [role, loc] if x)
                lines.append(r"\textbf{" + head + r"} \hfill " + right + r"\\")
                if sub:
                    lines.append(esc(sub) + r"\\")
                bl = e.get("bullets") or (
                    e.get("description") if isinstance(e.get("description"), list) else [])
                if bl:
                    lines.append(r"\begin{itemize}")
                    for b in bl:
                        lines.append(r"\item " + esc(b))
                    lines.append(r"\end{itemize}")
    lines.append(r"\end{document}")

    with open(tex_path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))
    return tex_path


def validate_resume(data):
    """数据完整性校验：返回 (errors, warnings, suggestions)。"""
    errors, warnings, suggestions = [], [], []
    pers = data.get("personal", {}) or {}
    name = str(pers.get("name", "")).strip()
    if not name or name == "姓名":
        errors.append("个人信息·姓名 缺失或仍为占位『姓名』")
    edu = data.get("education") or []
    if not any((e or {}).get("school") for e in edu):
        errors.append("教育经历·学校 缺失（至少填一段教育）")
    has_contact = any(str(pers.get(k, "")).strip() for k in ("phone", "email"))
    if not has_contact:
        warnings.append("联系方式（电话/邮箱）缺失，投递前必须补充")
    intern = data.get("internship") or []
    work = data.get("work") or []
    if not intern and not work:
        warnings.append("无实习/工作经历：建议用『模拟课题 / 模拟法庭 / 课程项目』替代，避免简历空洞")
        suggestions.append("替代方案见 resume_guide.md §10：① 模拟法庭（担任角色+书状/庭辩）；"
                           "② 课程/科研课题（负责检索与报告）；③ 法律援助/社团法务（文书/咨询）。")
    return errors, warnings, suggestions


def suggest_trim(data):
    """内容溢出时生成两份精简建议（保守 / 激进），返回 (保守文本, 激进文本)。"""
    optional = [k for k in ("research", "competition", "project", "summary")
                if data.get(k)]
    interns = data.get("internship") or []
    bullet_total = sum(len(i.get("bullets") or []) for i in interns)
    has_independent = any("独立" in b for i in interns for b in (i.get("bullets") or []))

    cons = ["【保守方案 · 保留全部模块，只压篇幅】"]
    if bullet_total > 6:
        cons.append(f"· 实习 bullet 偏多（共 {bullet_total} 条）：每段保留 2 条最强成果，"
                    f"把 3 条同类合并为 1 条概括（如『累计审阅合同约 X 份』）。")
    if has_independent:
        cons.append("· 将『独立完成』降级为『参与 / 协助』，减少篇幅也更稳妥。")
    cons.append("· 教育模块删掉非核心课程，仅留与求职方向最相关的 3 门。")
    cons.append("→ 适用：略超一页，想尽量保住经历。")

    aggr = ["【激进方案 · 优先砍可选模块】"]
    if optional:
        aggr.append("· 直接删除可选模块：" + "、".join(optional) + "（往往是溢出主因）。")
    else:
        aggr.append("· 当前无可选模块可删，从实习 bullet 入手。")
    aggr.append("· 实习每段压缩到 2 条 bullet；删除教育中的 courses / honors 次要行。")
    aggr.append("· 若写了个人总结，可整段删除。")
    aggr.append("→ 适用：远超一页，必须瘦身到一页。")

    return "\n".join(cons), "\n".join(aggr)


# ---------------------------------------------------------------------------
# 多版本：按版本 key 合并专属覆盖(versions_data)，并标注版本名
# ---------------------------------------------------------------------------
def build_version_data(data, vkey):
    d = dict(data)
    overrides = (data.get("versions_data") or {}).get(vkey) or {}
    for k, v in overrides.items():
        d[k] = v
    if vkey in VERSION_PRESETS:
        label = VERSION_PRESETS[vkey]["label"]
        pers = dict(d.get("personal", {}))
        ji = (pers.get("job_intention") or "")
        pers["job_intention"] = (ji + f"（{label}）") if ji else label
        d["personal"] = pers
    return d


def generate_one(data, cn_font, base_size_hint, T, accent_override,
                 out_dir, name, section_order=None, lang="zh", force_color=None):
    """生成单份(docx+pdf)，自动选字号保证一页；返回 (pdf_path, chosen, pages)。"""
    build_data, _ = resolve_lang_data(data, lang)
    docx_path = os.path.join(out_dir, f"{name}.docx")
    pdf_path = os.path.join(out_dir, f"{name}.pdf")
    candidate_sizes = [12, 11.5, 11, 10.5, 10, 9.5, 9, 8.5, 8, 7.5]
    chosen = None
    pages = None
    for i, base in enumerate(candidate_sizes):
        doc = build_document(build_data, cn_font, base, T, accent_override,
                             section_order, lang, force_color)
        doc.save(docx_path)
        ok, method = convert_to_pdf(docx_path, pdf_path)
        if not ok:
            return docx_path, None, None
        pages = count_pdf_pages(pdf_path)
        if pages <= 1:
            chosen = base
            break
        if i == len(candidate_sizes) - 1:
            chosen = base
    # 写入生成时间戳到 PDF 元数据
    stamp_pdf_metadata(pdf_path,
                       title=build_data.get("personal", {}).get("name", "Resume"))
    return docx_path, pdf_path, (chosen, pages)


# ---------------------------------------------------------------------------
# 入口
# ---------------------------------------------------------------------------
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--data", required=True, help="简历内容 JSON 路径")
    ap.add_argument("--out", required=True, help="输出目录")
    ap.add_argument("--font", default=None, help="中文字体: fang/yahei/song/kaiti")
    ap.add_argument("--template", default=None,
                    help="版式模板: classic/minimal/centered/modern/compact")
    ap.add_argument("--accent", default=None,
                    help="强调色: skyblue/black/gray/darkgray/navy/teal/... 或 #RRGGBB 自定义")
    ap.add_argument("--force-color", default=None, choices=["white", "black"],
                    help="强制标题文字颜色(覆盖自动反色, 不推荐): white/black")
    ap.add_argument("--style", default=None,
                    help="语言风格: aggressive(红圈所)/moderate(综合)/conservative(体制内)")
    ap.add_argument("--versions", default=None,
                    help="多版本简历: 逗号分隔 law/legal/civil，如 law,legal")
    ap.add_argument("--language", default=None,
                    help="语言: zh(中文)/en(英文)/both(中英各一份)")
    ap.add_argument("--format", default="both", dest="fmt",
                    help="输出格式: both(docx+pdf,默认)/docx/latex")
    ap.add_argument("--name", default="简历", help="输出文件名(不含扩展名)")
    args = ap.parse_args()
    logger.info(f"法学生求职简历生成器 v{__version__}")

    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)

    meta = data.get("meta", {})
    font_key = args.font or meta.get("font", "yahei")
    cn_font = CN_FONT_MAP.get(font_key, "微软雅黑")

    tpl_key = args.template or meta.get("template", "classic")
    T = TEMPLATES.get(tpl_key, TEMPLATES["classic"])
    accent_override = args.accent or meta.get("accent")

    # 强制标题文字颜色 (§13.2, 不推荐)
    force_color = {"white": WHITE, "black": DARK}.get(args.force_color) if args.force_color else None

    # 配色解析预检: 无效值给出友好报错, 而非抛出堆栈
    try:
        parse_color(accent_override or T.get("accent"))
    except ColorParseError as e:
        print(f"⛔ 配色解析失败：{e}")
        return
    logger.info(f"使用配色: {accent_override or T.get('accent')}")
    style = (args.style or meta.get("style") or "moderate").lower()
    if style not in ("aggressive", "moderate", "conservative"):
        style = "moderate"

    # 多版本解析
    versions = []
    if args.versions:
        versions = [v.strip().lower() for v in args.versions.split(",") if v.strip()]
    elif meta.get("versions"):
        versions = [v.strip().lower() for v in meta["versions"] if isinstance(v, str)]
    versions = [v for v in versions if v in VERSION_PRESETS]

    # 语言解析: zh / en / both
    lang_opt = (args.language or meta.get("language") or "zh").lower()
    if lang_opt == "both":
        languages = ["zh", "en"]
    elif lang_opt in ("zh", "en"):
        languages = [lang_opt]
    else:
        languages = ["zh"]

    # 输出格式
    fmt = (args.fmt or "both").lower()
    if fmt not in ("both", "docx", "latex"):
        fmt = "both"

    os.makedirs(args.out, exist_ok=True)

    # ---- 完整性校验 (每次生成前都做) ----
    errs, warns, sugg = validate_resume(data)
    if errs or warns:
        print("--- 完整性校验 ---")
        for e in errs:
            print(f"  [必填缺失] {e}")
        for w in warns:
            print(f"  [提醒] {w}")
        for s in sugg:
            print(f"  [建议] {s}")
    else:
        print("✅ 完整性校验通过")
    if errs:
        print("⛔ 存在必填缺失，已中止生成。请补齐后重试（见上）。")
        return

    style_label = {"aggressive": "红圈所·突出主导/结果",
                   "moderate": "综合·平衡团队与个人",
                   "conservative": "体制内·强调严谨/合规"}[style]
    print(f"🎨 语言风格：{style_label}")

    # ---- 组装生成单元 (支持多版本) ----
    units = []  # (data, name, section_order, version_label)
    if not versions:
        units.append((data, args.name, None, ""))
    else:
        print(f"📦 多版本生成（{len(versions)} 份）：" +
              "、".join(VERSION_PRESETS[v]["label"] for v in versions))
        for v in versions:
            units.append((build_version_data(data, v), f"{args.name}_{v}",
                          VERSION_PRESETS[v]["order"], VERSION_PRESETS[v]["label"]))

    lang_tag = {"zh": "中文", "en": "英文"}
    multi_lang = len(languages) > 1

    def lang_suffix(l):
        return "" if not multi_lang else ("_en" if l == "en" else "_zh")

    if multi_lang:
        print(f"🌐 语言输出：{'、'.join(lang_tag.get(l, l) + '版' for l in languages)}")
    if fmt == "latex":
        print("📐 输出格式：LaTeX (.tex，需自行用 XeLaTeX 编译)")

    for udata, uname, uorder, ulabel in units:
        for lang in languages:
            fname = uname + lang_suffix(lang)
            header_bits = []
            if ulabel:
                header_bits.append(f"版本：{ulabel}")
            if multi_lang:
                header_bits.append(f"{lang_tag.get(lang, lang)}版")
            if header_bits:
                print("\n=== " + " · ".join(header_bits) + " ===")
            # 英文版但缺 data['en'] 时提示 (不机器编造)
            if lang == "en":
                _, en_ok = resolve_lang_data(udata, "en")
                if not en_ok:
                    print("  [提醒] 未提供 data['en'] 英文内容块，英文版自由文本将回退中文；"
                          "建议补充英文 bullet 后重新生成（章节标题/联系标签已自动英译）。")
            if fmt == "latex":
                tex_path = os.path.join(args.out, f"{fname}.tex")
                export_latex(udata, tex_path, lang, uorder)
                print(f"✅ LaTeX 已导出: {tex_path}")
                print(f"   编译(中文需 ctex 宏包): xelatex {os.path.basename(tex_path)}")
            else:
                docx_path, pdf_path, res = generate_one(
                    udata, cn_font, 12, T, accent_override, args.out, fname,
                    uorder, lang, force_color)
                _report(tpl_key, accent_override, T, cn_font,
                        docx_path, pdf_path, res, udata)


def _report(tpl_key, accent_override, T, cn_font, docx_path, pdf_path, res, data):
    acc_name = (accent_override or T.get("accent"))
    chosen, pages = res
    if pdf_path is None:
        print("  [error] 无法生成 PDF (需要 Microsoft Word 或 LibreOffice)")
    elif pages and pages > 1:
        c, a = suggest_trim(data)
        print(f"\n⚠️ 内容较多, 已用最小字号 {chosen}pt 仍超出一页。两份精简建议：\n")
        print(c + "\n")
        print(a)
        print("\n请选择一份方案告诉我，或自行删减后重新生成。")
    else:
        print(f"✅ 一页简历已生成 (模板 {tpl_key}, 强调色 {acc_name}, "
              f"字号 {chosen}pt, 字体 {cn_font})")
    print(f"  DOCX: {docx_path}")
    if pdf_path:
        print(f"  PDF : {pdf_path}")


if __name__ == "__main__":
    try:
        main()
    except (ColorParseError, DataValidationError) as e:
        logger.error(f"生成失败: {e}")
        sys.exit(1)
    except Exception as e:
        logger.error(f"未知错误: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
